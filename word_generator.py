from docx import Document
from docx.shared import Pt, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from datetime import datetime

class WordGenerator:
    def __init__(self):
        self.document = Document()
        self._configure_document()
    
    def _configure_document(self):
        """Configura o estilo padrão do documento"""
        style = self.document.styles['Normal']
        style.font.name = 'Arial'
        style.font.size = Pt(12)
    
    def generate_prova(self, prova_data: dict) -> Document:
        """Gera um documento Word formatado com a prova"""
        # Cabeçalho
        header = self.document.add_paragraph()
        header.alignment = WD_ALIGN_PARAGRAPH.CENTER
        header.add_run("ESCOLA MODELO\n").bold = True
        header.add_run(f"Disciplina: {prova_data['disciplina']}\n")
        header.add_run(f"Série: {prova_data['serie']}\n")
        header.add_run(f"Professor(a): {prova_data['professor']}\n")
        header.add_run(f"Data: {prova_data['data']}\n\n")
        
        # Título
        title = self.document.add_paragraph()
        title.alignment = WD_ALIGN_PARAGRAPH.CENTER
        title.add_run(prova_data['titulo']).bold = True
        
        # Instruções
        self.document.add_paragraph("\nInstruções:")
        instrucoes = self.document.add_paragraph(
            "• Leia atentamente cada questão antes de respondê-la;\n"
            "• Cada questão tem apenas uma alternativa correta;\n"
            "• A prova deve ser feita individualmente e sem consulta;\n"
            "• Utilize caneta azul ou preta para marcar suas respostas."
        )
        instrucoes.paragraph_format.left_indent = Inches(0.5)
        
        # Questões
        self.document.add_paragraph("\nQuestões:")
        for questao in prova_data['questoes']:
            # Enunciado
            q_paragraph = self.document.add_paragraph()
            q_paragraph.add_run(f"{questao['numero']}. ").bold = True
            q_paragraph.add_run(questao['enunciado'])
            
            # Alternativas
            for letra, texto in questao['alternativas'].items():
                alt_paragraph = self.document.add_paragraph()
                alt_paragraph.paragraph_format.left_indent = Inches(0.5)
                alt_paragraph.add_run(f"{letra}) ").bold = True
                alt_paragraph.add_run(texto)
        
        # Gabarito (página separada)
        self.document.add_page_break()
        gabarito = self.document.add_paragraph()
        gabarito.alignment = WD_ALIGN_PARAGRAPH.CENTER
        gabarito.add_run("GABARITO\n\n").bold = True
        
        for questao in prova_data['questoes']:
            gab_line = self.document.add_paragraph()
            gab_line.alignment = WD_ALIGN_PARAGRAPH.CENTER
            gab_line.add_run(f"Questão {questao['numero']}: ").bold = True
            gab_line.add_run(f"Alternativa {questao['resposta_correta'].upper()}")
        
        return self.document
    
    def save_document(self, filepath: str):
        """Salva o documento Word no caminho especificado"""
        self.document.save(filepath)